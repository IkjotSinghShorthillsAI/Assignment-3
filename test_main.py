import os
import io
import csv
import tempfile
import unittest
from unittest.mock import patch, MagicMock, call
from PIL import Image

import pdfplumber
from docx.document import Document as DocxDocument

# Import the classes from your main.py
from main import (
    PDFLoader,
    DOCXLoader,
    PPTLoader,
    DataExtractor,
    FileStorage,
    SQLStorage,
)

# ----- Helpers for Fake Objects -----
# Fake PDF objects for extraction tests
class FakePDFPage:
    def __init__(self, blocks):
        self._blocks = blocks

    def get_text(self, mode):
        # mode is expected to be "dict"
        return {"blocks": self._blocks}

    def get_images(self, full=False):
        # Return a list with one dummy image reference.
        return [(10,)]  # tuple with a dummy xref

class FakePDFDoc:
    def __init__(self, pages):
        self.pages = pages

    def __len__(self):
        return len(self.pages)

    def __getitem__(self, index):
        return self.pages[index]

    def __iter__(self):
        return iter(self.pages)

    def extract_image(self, xref):
        # Create a small dummy image and return its PNG bytes.
        img = Image.new("RGB", (50, 50), color="blue")
        with io.BytesIO() as output:
            img.save(output, format="PNG")
            return {"image": output.getvalue()}

# ----- Test Cases for File Loaders -----
class TestFileLoaders(unittest.TestCase):
    def test_pdfloader_valid(self):
        loader = PDFLoader("sample.pdf")
        self.assertEqual(loader.file_path, "sample.pdf")
    
    def test_pdfloader_invalid(self):
        with self.assertRaises(ValueError):
            PDFLoader("sample.docx")
    
    def test_docxloader_valid(self):
        loader = DOCXLoader("sample.docx")
        self.assertEqual(loader.file_path, "sample.docx")
    
    def test_docxloader_invalid(self):
        with self.assertRaises(ValueError):
            DOCXLoader("sample.pdf")
    
    def test_pptloader_valid(self):
        loader = PPTLoader("sample.pptx")
        self.assertEqual(loader.file_path, "sample.pptx")
    
    def test_pptloader_invalid(self):
        with self.assertRaises(ValueError):
            PPTLoader("sample.docx")

# ----- Test Cases for DataExtractor -----
class TestDataExtractor(unittest.TestCase):
    def setUp(self):
        # Patch os.stat to return fixed metadata.
        self.fake_stat = os.stat_result((0, 0, 0, 0, 0, 0, 12345, 1600000000, 1600000001, 1600000002))
        patcher = patch("os.stat", return_value=self.fake_stat)
        self.addCleanup(patcher.stop)
        self.mock_stat = patcher.start()

    # A FakePDFLoader subclass to avoid opening real files.
    class FakePDFLoaderNoOpen(PDFLoader):
        def load_file(self):
            # Not used by extract_text or extract_images since we override fitts.open via patch.
            return None

    @patch("main.fitz.open")
    def test_extract_text_pdf(self, mock_fitz_open):
        # Create a fake PDF page with one text block.
        fake_blocks = [
            {
                "type": 0,
                "lines": [
                    {"spans": [{"text": "Hello PDF", "font": "BoldFont", "size": 14}]}
                ]
            }
        ]
        fake_page = FakePDFPage(fake_blocks)
        fake_doc = FakePDFDoc([fake_page])
        mock_fitz_open.return_value = fake_doc

        loader = self.FakePDFLoaderNoOpen("test.pdf")
        extractor = DataExtractor(loader)
        results = extractor.extract_text()
        expected = (1, "Hello PDF", "heading", "BoldFont", 14, True, False)
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0], expected)

    def test_extract_text_docx(self):
        # Create a fake DOCX document using MagicMock with spec=DocxDocument.
        fake_run = MagicMock()
        fake_run.text = "Hello DOCX"
        fake_run.font.name = "Regular"
        fake_run.font.size = MagicMock(pt=10)
        fake_run.bold = False
        fake_run.italic = False

        fake_para = MagicMock()
        fake_para.runs = [fake_run]
        # Ensure the paragraph text is available.
        fake_para.text = "Hello DOCX"

        fake_doc = MagicMock(spec=DocxDocument)
        fake_doc.paragraphs = [fake_para]

        # Create a fake DOCXLoader that returns our fake document.
        class FakeDOCXLoader(DOCXLoader):
            def load_file(self):
                return fake_doc

        loader = FakeDOCXLoader("test.docx")
        extractor = DataExtractor(loader)
        results = extractor.extract_text()
        expected = (1, "Hello DOCX", "text", "Regular", 10, False, False)
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0], expected)

    def test_extract_text_pptx(self):
        # Create a fake PPTX presentation with one slide and one text shape.
        # Simulate a shape with a text_frame containing one paragraph with one run.
        from collections import namedtuple
        FakePPTXRunObj = namedtuple("FakePPTXRunObj", ["text", "font", "hyperlink"])
        FakeFont = namedtuple("FakeFont", ["name", "size", "bold", "italic"])
        fake_run = FakePPTXRunObj("Hello PPTX", FakeFont("Regular", type("Size", (), {"pt": 10})(), False, False), None)
        fake_para = MagicMock()
        fake_para.runs = [fake_run]
        fake_text_frame = MagicMock()
        fake_text_frame.paragraphs = [fake_para]
        fake_shape = MagicMock()
        fake_shape.has_text_frame = True
        fake_shape.text_frame = fake_text_frame
        fake_slide = MagicMock()
        fake_slide.shapes = [fake_shape]
        fake_pptx = MagicMock()
        fake_pptx.slides = [fake_slide]

        class FakePPTLoader(PPTLoader):
            def load_file(self):
                return fake_pptx

        loader = FakePPTLoader("test.pptx")
        extractor = DataExtractor(loader)
        results = extractor.extract_text()
        expected = (1, "Hello PPTX", "text", "Regular", 10, False, False)
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0], expected)

    def test_extract_links_pdf(self):
        # Create a fake pdfplumber PDF object with an annotation.
        fake_annot = {"uri": "http://example.com"}
        fake_page = MagicMock()
        fake_page.annots = [fake_annot]
        fake_pdf = MagicMock(spec=pdfplumber.PDF)
        fake_pdf.pages = [fake_page]

        class FakePDFLoaderForLinks(PDFLoader):
            def load_file(self):
                return fake_pdf

        loader = FakePDFLoaderForLinks("test.pdf")
        extractor = DataExtractor(loader)
        links = extractor.extract_links()
        self.assertEqual(len(links), 1)
        self.assertEqual(links[0], (1, "http://example.com"))

    def test_extract_links_docx(self):
        # Create a fake DOCX document using MagicMock with spec=DocxDocument.
        fake_para = MagicMock()
        fake_para.text = "Check this link: http://example.com"
        fake_doc = MagicMock(spec=DocxDocument)
        fake_doc.paragraphs = [fake_para]
        fake_rel = MagicMock()
        fake_rel.reltype = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
        fake_rel.target_ref = "http://example.com"
        fake_doc.part = MagicMock()
        fake_doc.part.rels = {"rId1": fake_rel}

        class FakeDOCXLoader(DOCXLoader):
            def load_file(self):
                return fake_doc

        loader = FakeDOCXLoader("test.docx")
        extractor = DataExtractor(loader)
        links = extractor.extract_links()
        self.assertEqual(len(links), 1)
        self.assertEqual(links[0], (1, "http://example.com"))

    def test_extract_links_pptx(self):
        # Create a fake PPTX presentation with a hyperlink in a run.
        fake_hyperlink = MagicMock()
        fake_hyperlink.address = "http://example.com"
        fake_run = MagicMock()
        fake_run.text = "Link text"
        fake_run.font.name = "Regular"
        fake_run.font.size = MagicMock(pt=10)
        fake_run.font.bold = False
        fake_run.font.italic = False
        fake_run.hyperlink = fake_hyperlink
        fake_para = MagicMock()
        fake_para.runs = [fake_run]
        fake_text_frame = MagicMock()
        fake_text_frame.paragraphs = [fake_para]
        fake_shape = MagicMock()
        fake_shape.has_text_frame = True
        fake_shape.text_frame = fake_text_frame
        fake_slide = MagicMock()
        fake_slide.shapes = [fake_shape]
        fake_pptx = MagicMock()
        fake_pptx.slides = [fake_slide]

        class FakePPTLoader(PPTLoader):
            def load_file(self):
                return fake_pptx

        loader = FakePPTLoader("test.pptx")
        extractor = DataExtractor(loader)
        links = extractor.extract_links()
        self.assertEqual(len(links), 1)
        self.assertEqual(links[0], (1, "http://example.com"))

    @patch("main.fitz.open")
    def test_extract_images_pdf(self, mock_fitz_open):
        # Setup fake PDF pages with images.
        fake_page = FakePDFPage([])  # No text blocks needed.
        fake_doc = FakePDFDoc([fake_page])
        mock_fitz_open.return_value = fake_doc

        # Simulate that the page returns one image.
        fake_page.get_images = MagicMock(return_value=[(10,)])
        # Generate PNG bytes for the fake image.
        fake_img = Image.new("RGB", (60, 60))
        with io.BytesIO() as output:
            fake_img.save(output, format="PNG")
            image_bytes = output.getvalue()
        fake_doc.extract_image = MagicMock(return_value={"image": image_bytes})

        loader = self.FakePDFLoaderNoOpen("test.pdf")
        extractor = DataExtractor(loader)
        images = extractor.extract_images()
        self.assertEqual(len(images), 1)
        page_num, img_format, size, img_obj = images[0]
        self.assertEqual(page_num, 1)
        self.assertEqual(img_format, "PNG")
        self.assertEqual(size, (60, 60))

    def test_extract_images_docx(self):
        # For DOCX, simulate an image relationship.
        img = Image.new("RGB", (70, 70))
        with io.BytesIO() as output:
            img.save(output, format="PNG")
            image_bytes = output.getvalue()
        fake_rel = MagicMock()
        fake_rel.target_ref = "word/media/image1.png"
        fake_rel.target_part = MagicMock(blob=image_bytes)
        fake_doc = MagicMock(spec=DocxDocument)
        fake_doc.part = MagicMock()
        fake_doc.part.rels = {"rId1": fake_rel}
        fake_doc.paragraphs = []  # paragraphs not needed for images.

        class FakeDOCXLoader(DOCXLoader):
            def load_file(self):
                return fake_doc

        loader = FakeDOCXLoader("test.docx")
        extractor = DataExtractor(loader)
        images = extractor.extract_images()
        self.assertEqual(len(images), 1)
        page_num, img_format, size, img_obj = images[0]
        self.assertEqual(page_num, 1)
        self.assertEqual(img_format, "PNG")
        self.assertEqual(size, img.size)

    def test_extract_images_pptx(self):
        # For PPTX, simulate a slide with a picture shape.
        img = Image.new("RGB", (80, 80))
        with io.BytesIO() as output:
            img.save(output, format="PNG")
            image_bytes = output.getvalue()

        fake_shape = MagicMock()
        fake_shape.shape_type = 13  # Indicates a picture.
        fake_image = MagicMock()
        fake_image.blob = image_bytes
        fake_shape.image = fake_image
        fake_slide = MagicMock()
        fake_slide.shapes = [fake_shape]
        fake_pptx = MagicMock()
        fake_pptx.slides = [fake_slide]

        class FakePPTLoader(PPTLoader):
            def load_file(self):
                return fake_pptx

        loader = FakePPTLoader("test.pptx")
        extractor = DataExtractor(loader)
        images = extractor.extract_images()
        self.assertEqual(len(images), 1)
        page_num, img_format, size, img_obj = images[0]
        self.assertEqual(page_num, 1)
        self.assertEqual(img_format, "PNG")
        self.assertEqual(size, img.size)

    def test_extract_tables_pdf(self):
        # For PDF, simulate a page that returns tables.
        fake_table = [["cell1", "cell2"], ["cell3", "cell4"]]
        fake_page = MagicMock()
        fake_page.extract_tables = MagicMock(return_value=[fake_table])
        fake_pdf = MagicMock(spec=pdfplumber.PDF)
        fake_pdf.pages = [fake_page]

        class FakePDFLoaderForTables(PDFLoader):
            def load_file(self):
                return fake_pdf

        loader = FakePDFLoaderForTables("test.pdf")
        extractor = DataExtractor(loader)
        tables = extractor.extract_tables()
        self.assertEqual(len(tables), 1)
        page_num, num_rows, num_cols, table = tables[0]
        self.assertEqual(page_num, 1)
        self.assertEqual(num_rows, 2)
        self.assertEqual(num_cols, 2)
        self.assertEqual(table, fake_table)

    def test_extract_tables_docx(self):
        # For DOCX, simulate a document with one table using a MagicMock.
        # Create fake cell objects.
        fake_cell1 = MagicMock()
        fake_cell1.text = "A"
        fake_cell2 = MagicMock()
        fake_cell2.text = "B"
        fake_row1 = MagicMock()
        fake_row1.cells = [fake_cell1, fake_cell2]
        fake_cell3 = MagicMock()
        fake_cell3.text = "C"
        fake_cell4 = MagicMock()
        fake_cell4.text = "D"
        fake_row2 = MagicMock()
        fake_row2.cells = [fake_cell3, fake_cell4]
        fake_table = MagicMock()
        fake_table.rows = [fake_row1, fake_row2]

        fake_doc = MagicMock(spec=DocxDocument)
        fake_doc.tables = [fake_table]
        fake_doc.paragraphs = []  # paragraphs not needed.

        class FakeDOCXLoader(DOCXLoader):
            def load_file(self):
                return fake_doc

        loader = FakeDOCXLoader("test.docx")
        extractor = DataExtractor(loader)
        tables = extractor.extract_tables()
        self.assertEqual(len(tables), 1)
        page_num, num_rows, num_cols, table = tables[0]
        self.assertEqual(num_rows, 2)
        self.assertEqual(num_cols, 2)
        self.assertEqual(table, [["A", "B"], ["C", "D"]])

    def test_extract_tables_pptx(self):
        # For PPTX, simulate a slide with a shape that has a table.
        fake_shape = MagicMock()
        fake_shape.has_table = True
        # Create a fake table with 2 rows and 2 cells each.
        FakeCell = lambda t: type("FakeCell", (), {"text": t})()
        FakeRow = lambda cells: type("FakeRow", (), {"cells": cells})()
        fake_table_obj = type("FakeTable", (), {"rows": [FakeRow([FakeCell("cell"), FakeCell("cell")]),
                                                         FakeRow([FakeCell("cell"), FakeCell("cell")])]})
        fake_shape.table = fake_table_obj
        fake_slide = MagicMock()
        fake_slide.shapes = [fake_shape]
        fake_pptx = MagicMock()
        fake_pptx.slides = [fake_slide]

        class FakePPTLoader(PPTLoader):
            def load_file(self):
                return fake_pptx

        loader = FakePPTLoader("test.pptx")
        extractor = DataExtractor(loader)
        tables = extractor.extract_tables()
        self.assertEqual(len(tables), 1)
        page_num, num_rows, num_cols, table = tables[0]
        self.assertEqual(num_rows, 2)
        self.assertEqual(num_cols, 2)

# ----- Test Cases for Storage Classes -----
class TestFileStorage(unittest.TestCase):
    def setUp(self):
        self.fake_extractor = MagicMock()
        self.fake_extractor.extract_text.return_value = [
            (1, "Sample text", "text", "Font", 10, False, False)
        ]
        self.fake_extractor.extract_tables.return_value = [
            (1, 1, 2, [["cell1", "cell2"]])
        ]
        img = Image.new("RGB", (100, 100))
        self.fake_extractor.extract_images.return_value = [
            (1, "PNG", img.size, img)
        ]
        self.fake_extractor.metadata = {"file_size": 12345, "creation_time": 1600000000, "modification_time": 1600000001}
        self.fake_extractor.file_name = "dummy_file"

    def test_file_storage_save_data(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            storage = FileStorage(self.fake_extractor)
            storage.save_data(tmpdir)
            text_file = os.path.join(tmpdir, "text_data.txt")
            self.assertTrue(os.path.exists(text_file))
            with open(text_file, "r", encoding="utf-8") as f:
                content = f.read()
                self.assertIn("Sample text", content)
            tables_file = os.path.join(tmpdir, "tables.csv")
            self.assertTrue(os.path.exists(tables_file))
            with open(tables_file, "r", encoding="utf-8") as f:
                reader = csv.reader(f)
                rows = list(reader)
                self.assertTrue(any("cell1" in cell for row in rows for cell in row))
            image_path = os.path.join(tmpdir, "image_1.png")
            self.assertTrue(os.path.exists(image_path))

class TestSQLStorage(unittest.TestCase):
    def setUp(self):
        self.fake_extractor = MagicMock()
        self.fake_extractor.extract_text.return_value = [
            (1, "SQL text", "text", "Font", 12, True, False)
        ]
        self.fake_extractor.extract_links.return_value = [
            (1, "http://sqltest.com")
        ]
        self.fake_extractor.extract_tables.return_value = [
            (1, 2, 2, [["A", "B"], ["C", "D"]])
        ]
        self.fake_extractor.metadata = {"file_size": 54321, "creation_time": 1600100000, "modification_time": 1600100001}
        self.fake_extractor.file_name = "sql_dummy_file"

    @patch("main.mysql.connector.connect")
    def test_sql_storage_save_data(self, mock_connect):
        fake_cursor = MagicMock()
        fake_conn = MagicMock()
        fake_conn.cursor.return_value = fake_cursor
        mock_connect.return_value = fake_conn

        storage = SQLStorage(self.fake_extractor, host="localhost", user="root", password="pass", database="test_db")
        storage.save_data()

        create_db_call = call("CREATE DATABASE IF NOT EXISTS test_db")
        self.assertIn(create_db_call, fake_cursor.execute.call_args_list)

        insert_calls = [args for args in fake_cursor.execute.call_args_list if "INSERT INTO extracted_data" in args[0][0]]
        self.assertEqual(len(insert_calls), 3)

        fake_conn.commit.assert_called_once()
        # Expect two calls to close: one from _ensure_database_exists and one from save_data.
        self.assertEqual(fake_conn.close.call_count, 2)

# ----- Run all tests -----
if __name__ == "__main__":
    unittest.main()
